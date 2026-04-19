"""
Advanced Features Module for Ariadne.

Provides:
- Intelligent summarization with multi-language support
- Knowledge graph visualization (DOT, JSON, interactive HTML)
- Export to multiple formats (Markdown, HTML, Word, PDF)
- Reranking and semantic search enhancement
"""

from __future__ import annotations

import json
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass

from ariadne.ingest.base import Document
from ariadne.graph.models import Entity, Relation, GraphDocument
from ariadne.graph.storage import GraphStorage
from ariadne.llm.base import BaseLLM
from ariadne.config import get_config, AriadneConfig


# ============================================================================
# Language Support
# ============================================================================

LANGUAGE_PROMPTS = {
    "zh_CN": {
        "summarize": "请用简体中文总结以下内容：",
        "entities": "请用简体中文列出实体名称：",
        "relations": "请用简体中文描述关系：",
    },
    "zh_TW": {
        "summarize": "请用繁体中文总结以下内容：",
        "entities": "请用繁体中文列出实体名称：",
        "relations": "请用繁体中文描述关系：",
    },
    "en": {
        "summarize": "Please summarize the following in English:",
        "entities": "Please list entity names in English:",
        "relations": "Please describe relationships in English:",
    },
    "fr": {
        "summarize": "Veuillez resumer le texte suivant en francais :",
        "entities": "Veuillez lister les noms d'entites en francais :",
        "relations": "Veuillez decrire les relations en francais :",
    },
    "es": {
        "summarize": "Por favor, resuma el siguiente texto en espanol:",
        "entities": "Por favor, liste los nombres de entidades en espanol:",
        "relations": "Por favor, describa las relaciones en espanol:",
    },
    "ru": {
        "summarize": "Пожалуйста, суммируйте следующий текст на русском языке:",
        "entities": "Пожалуйста, перечислите названия сущностей на русском языке:",
        "relations": "Пожалуйста, опишите отношения на русском языке:",
    },
    "ar": {
        "summarize": "yruhafat alnass althaqyfy bialrbyt alarbyt:",
        "entities": "yrfy tskhys asmaw alkyanan bialrbyt alarbyt:",
        "relations": "yrfy wsrf alarabat bialrbyt alarbyt:",
    },
}


def get_language_prompt(language: str, prompt_type: str = "summarize") -> str:
    """Get a localized prompt for the specified language."""
    prompts = LANGUAGE_PROMPTS.get(language, LANGUAGE_PROMPTS["en"])
    return prompts.get(prompt_type, prompts["summarize"])


# ============================================================================
# Summarizer
# ============================================================================

@dataclass
class SummaryResult:
    """Result of a summarization operation."""
    summary: str
    keywords: List[str]
    topics: List[str]
    language: str
    sources: List[str]


class Summarizer:
    """
    Intelligent document summarizer with multi-language support.
    
    Features:
    - LLM-powered summarization
    - Multi-language output (7 UN languages supported)
    - Keyword and topic extraction
    - Source attribution
    
    Example:
        summarizer = Summarizer(llm)
        result = summarizer.summarize(
            documents=[doc1, doc2],
            language="zh_CN"
        )
    """
    
    SUMMARY_PROMPT = """You are an expert at summarizing documents. Given the following content, provide a comprehensive summary.

Requirements:
1. Main topics and themes (3-5 key topics)
2. Key insights or findings (bullet points)
3. A concise summary paragraph (2-3 sentences)
4. Important keywords (5-10 keywords)

Format your response as JSON:
{{
    "summary": "A concise 2-3 sentence summary paragraph",
    "topics": ["topic1", "topic2", "topic3"],
    "keywords": ["keyword1", "keyword2", "keyword3", "keyword4", "keyword5"],
    "insights": [
        {{
            "text": "Key insight from the documents",
            "source": "document name or path"
        }}
    ]
}}

Content to summarize:
{content}
"""
    
    def __init__(
        self,
        llm: Optional[BaseLLM] = None,
        config: Optional[AriadneConfig] = None,
    ):
        """
        Initialize the summarizer.
        
        Args:
            llm: LLM instance for summarization. If None, uses config.
            config: Configuration. If None, auto-loads from config.
        """
        self.llm = llm
        self.config = config or get_config()
        
        if llm is None:
            self.llm = self.config.create_llm()
    
    def summarize(
        self,
        documents: List[Document],
        query: Optional[str] = None,
        language: Optional[str] = None,
    ) -> SummaryResult:
        """
        Generate a summary of documents.
        
        Args:
            documents: List of documents to summarize.
            query: Optional search query for targeted summarization.
            language: Output language (defaults to config setting).
            
        Returns:
            SummaryResult with summary, keywords, and topics.
        """
        if not documents:
            return SummaryResult(
                summary="No documents to summarize.",
                keywords=[],
                topics=[],
                language=language or self.config.get_output_language(),
                sources=[],
            )
        
        if self.llm is None:
            return self._basic_summary(documents, language)
        
        language = language or self.config.get_output_language()
        lang_prompt = get_language_prompt(language, "summarize")
        
        # Prepare content
        content_parts = []
        sources = []
        
        for i, doc in enumerate(documents):
            sources.append(doc.source_path or f"Document {i+1}")
            content_parts.append(f"[{sources[-1]}]\n{doc.content[:2000]}")
        
        content = "\n\n".join(content_parts)
        
        if query:
            content = f"Query: {query}\n\n{content}"
        
        # Build prompt — inject language instruction at the TOP so it dominates
        prompt = self.SUMMARY_PROMPT.format(content=content)
        # Prepend the localized language instruction; this tells the LLM what
        # language to use for its entire response.
        prompt = lang_prompt + "\n\n" + prompt
        
        try:
            response = self.llm.chat(prompt)
            
            # Parse response
            result = self._parse_summary_response(response.content, language, sources)
            return result
            
        except Exception as e:
            return SummaryResult(
                summary=f"Error generating summary: {str(e)}",
                keywords=[],
                topics=[],
                language=language,
                sources=sources,
            )
    
    def _parse_summary_response(
        self,
        content: str,
        language: str,
        sources: List[str],
    ) -> SummaryResult:
        """Parse LLM response into SummaryResult."""
        try:
            # Try to extract JSON
            json_str = self._extract_json(content)
            data = json.loads(json_str)

            # Ensure data is a dict (LLM may occasionally return a JSON string)
            if not isinstance(data, dict):
                raise json.JSONDecodeError("Expected JSON object", json_str, 0)

            summary = data.get("summary", "")
            keywords = data.get("keywords", [])
            topics = data.get("topics", [])

            # Coerce to expected types if LLM returned wrong shapes
            if not isinstance(summary, str):
                summary = str(summary)
            if not isinstance(keywords, list):
                keywords = []
            if not isinstance(topics, list):
                topics = []

            return SummaryResult(
                summary=summary,
                keywords=keywords,
                topics=topics,
                language=language,
                sources=sources,
            )
        except (json.JSONDecodeError, ValueError):
            # Fall back to plain text — strip stray JSON punctuation for readability
            clean = content.strip().lstrip("{").rstrip("}")
            return SummaryResult(
                summary=clean[:1000],
                keywords=[],
                topics=[],
                language=language,
                sources=sources,
            )

    def _extract_json(self, content: str) -> str:
        """Extract JSON from content, handling markdown code blocks."""
        content = content.strip()

        # Handle ```json ... ``` blocks
        if "```json" in content:
            parts = content.split("```json", 1)
            if len(parts) > 1:
                content = parts[1].split("```")[0]
        # Handle ``` ... ``` blocks (may have a language tag on the first line)
        elif "```" in content:
            parts = content.split("```", 1)
            if len(parts) > 1:
                inner = parts[1].split("```")[0]
                # Drop leading language tag line (e.g. "json\n{...}")
                lines = inner.splitlines()
                if lines and not lines[0].strip().startswith("{"):
                    inner = "\n".join(lines[1:])
                content = inner

        # If no code block found, try to locate the first { ... } in the response
        stripped = content.strip()
        if not stripped.startswith("{"):
            start = stripped.find("{")
            end = stripped.rfind("}")
            if start != -1 and end != -1 and end > start:
                stripped = stripped[start:end + 1]
            content = stripped

        return content.strip()
    
    def _basic_summary(
        self,
        documents: List[Document],
        language: Optional[str],
    ) -> SummaryResult:
        """Basic summary without LLM."""
        sources = [doc.source_path or f"Doc {i+1}" for i, doc in enumerate(documents)]
        
        # Extract some keywords from content
        words = set()
        for doc in documents:
            content_words = doc.content.split()[:200]
            words.update(content_words)
        
        return SummaryResult(
            summary=f"Found {len(documents)} documents containing approximately {sum(len(d.content.split()) for d in documents)} words.",
            keywords=list(words)[:10],
            topics=[],
            language=language or "en",
            sources=sources,
        )


# ============================================================================
# Graph Visualizer
# ============================================================================

class GraphVisualizer:
    """
    Knowledge graph visualization exporter.
    
    Supports multiple output formats:
    - DOT: GraphViz format
    - JSON: For web visualizations
    - HTML: Interactive visualization
    - Mermaid: Markdown-compatible diagrams
    
    Example:
        visualizer = GraphVisualizer(graph_storage)
        dot_output = visualizer.to_dot()
        html_output = visualizer.to_html()
    """
    
    def __init__(
        self,
        graph: GraphStorage,
        config: Optional[AriadneConfig] = None,
    ):
        """
        Initialize the visualizer.
        
        Args:
            graph: GraphStorage instance.
            config: Configuration for styling.
        """
        self.graph = graph
        self.config = config or get_config()
    
    def to_dot(
        self,
        entity_ids: Optional[List[str]] = None,
        max_nodes: int = 50,
    ) -> str:
        """
        Export to DOT format (GraphViz).
        
        Args:
            entity_ids: Optional list of entity IDs to include. If None, uses all.
            max_nodes: Maximum number of nodes to include.
            
        Returns:
            DOT format string.
        """
        lines = [
            "digraph KnowledgeGraph {",
            '    rankdir=LR;',
            '    node [shape=box, style="rounded,filled"];',
            '    edge [arrowhead=normal];',
            "",
        ]
        
        # Color mapping for entity types
        type_colors = {
            "person": "#E3F2FD",
            "organization": "#FFF3E0",
            "location": "#E8F5E9",
            "event": "#FCE4EC",
            "concept": "#F3E5F5",
            "work": "#FFFDE7",
            "topic": "#ECEFF1",
            "unknown": "#FAFAFA",
        }
        
        # Get entities to include
        if entity_ids:
            entities = [self.graph.get_entity(eid) for eid in entity_ids]
            entities = [e for e in entities if e is not None]
        else:
            entities = self.graph.get_all_entities()[:max_nodes]
        
        entity_ids = {e.entity_id for e in entities}
        
        # Add nodes
        for entity in entities:
            color = type_colors.get(entity.entity_type.value, type_colors["unknown"])
            label = entity.name.replace('"', '\\"')
            if entity.description:
                label = f'{label}\\n{entity.description[:50]}'
            
            lines.append(
                f'    "{entity.entity_id}" [label="{label}", '
                f'fillcolor="{color}", tooltip="{entity.entity_type.value}"];'
            )
        
        lines.append("")
        
        # Add edges (relations)
        for relation in self.graph.get_all_relations():
            if relation.source_id in entity_ids and relation.target_id in entity_ids:
                rel_type = relation.relation_type.value.replace("_", " ")
                lines.append(
                    f'    "{relation.source_id}" -> "{relation.target_id}" '
                    f'[label="{rel_type}", tooltip="{relation.description}"];'
                )
        
        lines.append("}")
        
        return "\n".join(lines)
    
    def to_json(
        self,
        entity_ids: Optional[List[str]] = None,
        max_nodes: int = 100,
    ) -> str:
        """
        Export to JSON format for web visualizations.
        
        Returns:
            JSON string with nodes and edges.
        """
        nodes = []
        edges = []
        
        # Get entities
        if entity_ids:
            entities = [self.graph.get_entity(eid) for eid in entity_ids]
            entities = [e for e in entities if e is not None]
        else:
            entities = self.graph.get_all_entities()[:max_nodes]
        
        entity_ids = {e.entity_id for e in entities}
        
        # Build nodes
        for entity in entities:
            nodes.append({
                "id": entity.entity_id,
                "label": entity.name,
                "type": entity.entity_type.value,
                "description": entity.description,
                "aliases": entity.aliases,
            })
        
        # Build edges
        for relation in self.graph.get_all_relations():
            if relation.source_id in entity_ids and relation.target_id in entity_ids:
                edges.append({
                    "id": relation.relation_id,
                    "source": relation.source_id,
                    "target": relation.target_id,
                    "type": relation.relation_type.value,
                    "description": relation.description,
                })
        
        return json.dumps({
            "nodes": nodes,
            "edges": edges,
            "stats": self.graph.stats(),
        }, indent=2, ensure_ascii=False)
    
    def to_mermaid(
        self,
        entity_ids: Optional[List[str]] = None,
        max_nodes: int = 30,
    ) -> str:
        """
        Export to Mermaid diagram format.
        
        Returns:
            Mermaid flowchart string.
        """
        lines = ["flowchart TD"]
        
        # Get entities
        if entity_ids:
            entities = [self.graph.get_entity(eid) for eid in entity_ids]
            entities = [e for e in entities if e is not None]
        else:
            entities = self.graph.get_all_entities()[:max_nodes]
        
        entity_ids = {e.entity_id for e in entities}
        
        # Add nodes with style classes
        for entity in entities:
            safe_id = entity.entity_id[:8]
            safe_name = entity.name.replace('"', "'")[:30]
            lines.append(f'    {safe_id}["{safe_name}"]')
        
        lines.append("")
        
        # Add edges
        for relation in self.graph.get_all_relations():
            if relation.source_id in entity_ids and relation.target_id in entity_ids:
                src = relation.source_id[:8]
                tgt = relation.target_id[:8]
                rel = relation.relation_type.value[:20]
                lines.append(f'    {src} -->|{rel}| {tgt}')
        
        return "\n".join(lines)
    
    def to_html(
        self,
        entity_ids: Optional[List[str]] = None,
        max_nodes: int = 50,
        title: str = "Knowledge Graph",
    ) -> str:
        """
        Generate an interactive HTML visualization.
        
        Uses vis.js for interactive graph display.
        
        Returns:
            Complete HTML string.
        """
        # Get data
        if entity_ids:
            entities = [self.graph.get_entity(eid) for eid in entity_ids]
            entities = [e for e in entities if e is not None]
        else:
            entities = self.graph.get_all_entities()[:max_nodes]
        
        entity_ids = {e.entity_id for e in entities}
        
        # Build nodes and edges for vis.js
        nodes = []
        edges = []
        
        type_colors = {
            "person": "#4CAF50",
            "organization": "#2196F3",
            "location": "#FF9800",
            "event": "#E91E63",
            "concept": "#9C27B0",
            "work": "#FFC107",
            "topic": "#607D8B",
            "unknown": "#9E9E9E",
        }
        
        for entity in entities:
            color = type_colors.get(entity.entity_type.value, type_colors["unknown"])
            nodes.append({
                "id": entity.entity_id[:16],
                "label": entity.name[:30],
                "title": f"{entity.entity_type.value}: {entity.description[:100]}",
                "color": color,
                "font": {"color": "#333333"},
            })
        
        for relation in self.graph.get_all_relations():
            if relation.source_id in entity_ids and relation.target_id in entity_ids:
                edges.append({
                    "from": relation.source_id[:16],
                    "to": relation.target_id[:16],
                    "label": relation.relation_type.value[:15],
                    "arrows": "to",
                })
        
        nodes_json = json.dumps(nodes, ensure_ascii=False)
        edges_json = json.dumps(edges, ensure_ascii=False)
        
        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <script src="https://unpkg.com/vis-network/standalone/umd/vis-network.min.js"></script>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            margin: 0;
            padding: 20px;
            background: #f5f5f5;
        }}
        h1 {{
            color: #333;
            text-align: center;
        }}
        #network {{
            width: 100%;
            height: 600px;
            border: 1px solid #ddd;
            background: white;
            border-radius: 8px;
        }}
        .legend {{
            display: flex;
            flex-wrap: wrap;
            gap: 15px;
            justify-content: center;
            margin: 20px 0;
        }}
        .legend-item {{
            display: flex;
            align-items: center;
            gap: 5px;
        }}
        .legend-color {{
            width: 16px;
            height: 16px;
            border-radius: 50%;
        }}
        .stats {{
            text-align: center;
            color: #666;
            margin-top: 10px;
        }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    
    <div class="legend">
        <div class="legend-item"><div class="legend-color" style="background:#4CAF50"></div>Person</div>
        <div class="legend-item"><div class="legend-color" style="background:#2196F3"></div>Organization</div>
        <div class="legend-item"><div class="legend-color" style="background:#FF9800"></div>Location</div>
        <div class="legend-item"><div class="legend-color" style="background:#E91E63"></div>Event</div>
        <div class="legend-item"><div class="legend-color" style="background:#9C27B0"></div>Concept</div>
        <div class="legend-item"><div class="legend-color" style="background:#FFC107"></div>Work</div>
        <div class="legend-item"><div class="legend-color" style="background:#607D8B"></div>Topic</div>
    </div>
    
    <div id="network"></div>
    
    <p class="stats">
        Nodes: {len(nodes)} | Edges: {len(edges)} |
        Total entities: {self.graph.stats()['num_entities']}
    </p>
    
    <script>
        var nodes = new vis.DataSet({nodes_json});
        var edges = new vis.DataSet({edges_json});
        
        var container = document.getElementById('network');
        var data = {{ nodes: nodes, edges: edges }};
        var options = {{
            nodes: {{
                shape: 'dot',
                size: 16,
                borderWidth: 2,
                shadow: true
            }},
            edges: {{
                width: 1,
                smooth: {{ type: 'continuous' }}
            }},
            physics: {{
                stabilization: {{ iterations: 200 }},
                barnesHut: {{
                    gravitationalConstant: -2000,
                    springConstant: 0.04,
                    springLength: 150
                }}
            }},
            interaction: {{
                hover: true,
                tooltipDelay: 200
            }}
        }};
        
        var network = new vis.Network(container, data, options);
    </script>
</body>
</html>"""
        
        return html


# ============================================================================
# Exporter
# ============================================================================

class Exporter:
    """
    Export Ariadne data to multiple formats.
    
    Supports:
    - Markdown (default, portable)
    - HTML (formatted, interactive)
    - Word (.docx, requires python-docx)
    - PDF (requires reportlab or weasyprint)
    
    Example:
        exporter = Exporter(graph_storage)
        exporter.export("output.html", format="html")
    """
    
    def __init__(
        self,
        graph: Optional[GraphStorage] = None,
        config: Optional[AriadneConfig] = None,
    ):
        """
        Initialize the exporter.
        
        Args:
            graph: Optional GraphStorage for knowledge graph export.
            config: Configuration for formatting options.
        """
        self.graph = graph
        self.config = config or get_config()
    
    def export(
        self,
        output_path: str,
        format: str = "markdown",
        title: str = "Ariadne Export",
        include_graph: bool = True,
        language: Optional[str] = None,
    ) -> str:
        """
        Export to specified format.
        
        Args:
            output_path: Output file path.
            format: Format (markdown, html, docx, pdf).
            title: Document title.
            include_graph: Include knowledge graph section.
            language: Output language.
            
        Returns:
            Path to output file.
        """
        language = language or self.config.get_output_language()
        
        if format == "markdown":
            content = self._to_markdown(title, include_graph, language)
        elif format == "html":
            content = self._to_html(title, include_graph, language)
        elif format == "docx":
            return self._to_docx(title, include_graph, language, output_path)
        elif format == "pdf":
            return self._to_pdf(title, include_graph, language, output_path)
        else:
            raise ValueError(f"Unsupported format: {format}")
        
        # Write content to file
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
        
        return output_path
    
    def _to_markdown(
        self,
        title: str,
        include_graph: bool,
        language: str,
    ) -> str:
        """Generate Markdown export."""
        lines = [
            f"# {title}",
            "",
            f"*Generated by Ariadne - {language}*",
            "",
        ]
        
        # Graph section
        if include_graph and self.graph:
            stats = self.graph.stats()
            
            lines.extend([
                "## Knowledge Graph",
                "",
                f"- **Entities**: {stats['num_entities']}",
                f"- **Relations**: {stats['num_relations']}",
                f"- **Entity Types**: {stats['num_entity_types']}",
                "",
            ])
            
            # List entities by type
            lines.append("### Entities")
            lines.append("")
            
            for entity_type in ["person", "organization", "location", "event", "concept", "work"]:
                entities = self.graph.get_entities_by_type(entity_type)
                if entities:
                    lines.append(f"#### {entity_type.title()}")
                    for entity in entities[:20]:
                        lines.append(f"- **{entity.name}**: {entity.description[:100]}")
                    lines.append("")
            
            # List relations
            lines.append("### Relations")
            lines.append("")
            
            relations = self.graph.get_all_relations()[:30]
            for relation in relations:
                src = self.graph.get_entity(relation.source_id)
                tgt = self.graph.get_entity(relation.target_id)
                if src and tgt:
                    lines.append(
                        f"- {src.name} **({relation.relation_type.value})** {tgt.name}"
                    )
            
            lines.append("")
        
        lines.append("---")
        lines.append("*Export generated by Ariadne*")
        
        return "\n".join(lines)
    
    def _to_html(
        self,
        title: str,
        include_graph: bool,
        language: str,
    ) -> str:
        """Generate HTML export with optional embedded graph."""
        graph_html = ""
        
        if include_graph and self.graph:
            visualizer = GraphVisualizer(self.graph, self.config)
            graph_html = f"""
            <section>
                <h2>Knowledge Graph</h2>
                {visualizer.to_html(title="Interactive Graph", max_nodes=30)}
            </section>
            """
        
        html = f"""<!DOCTYPE html>
<html lang="{language}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{ color: #1a1a2e; border-bottom: 3px solid #4CAF50; padding-bottom: 10px; }}
        h2 {{ color: #16213e; margin-top: 40px; }}
        h3 {{ color: #0f3460; }}
        .metadata {{ color: #666; font-style: italic; margin-bottom: 30px; }}
        .entity {{ margin: 10px 0; padding: 10px; background: #f8f9fa; border-radius: 5px; }}
        .relation {{ margin: 5px 0; padding-left: 20px; color: #555; }}
        footer {{ margin-top: 50px; padding-top: 20px; border-top: 1px solid #ddd; color: #888; text-align: center; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p class="metadata">Generated by Ariadne | Language: {language}</p>
    
    {graph_html}
    
    <footer>
        <p>Generated by Ariadne Memory System</p>
    </footer>
</body>
</html>"""
        
        return html
    
    def _to_docx(
        self,
        title: str,
        include_graph: bool,
        language: str,
        output_path: str,
    ) -> str:
        """Generate Word document."""
        try:
            from docx import Document
            
            doc = Document()
            doc.add_heading(title, 0)
            doc.add_paragraph(f"Language: {language}")
            
            if include_graph and self.graph:
                doc.add_heading("Knowledge Graph", 1)
                
                stats = self.graph.stats()
                doc.add_paragraph(f"Total Entities: {stats['num_entities']}")
                doc.add_paragraph(f"Total Relations: {stats['num_relations']}")
                
                # Add entities
                doc.add_heading("Entities", 2)
                for entity_type in ["person", "organization", "location"]:
                    entities = self.graph.get_entities_by_type(entity_type)
                    if entities:
                        doc.add_heading(entity_type.title(), 3)
                        for entity in entities[:10]:
                            doc.add_paragraph(f"{entity.name}: {entity.description[:100]}")
            
            doc.save(output_path)
            return output_path
            
        except ImportError:
            # Fallback to markdown
            content = self._to_markdown(title, include_graph, language)
            with open(output_path.replace(".docx", ".md"), "w", encoding="utf-8") as f:
                f.write(content)
            return output_path.replace(".docx", ".md")
    
    def _to_pdf(
        self,
        title: str,
        include_graph: bool,
        language: str,
        output_path: str,
    ) -> str:
        """Generate PDF document."""
        try:
            from weasyprint import HTML as WeasyHTML
            
            html = self._to_html(title, include_graph, language)
            WeasyHTML(string=html).write_pdf(output_path)
            return output_path
        except ImportError:
            # Fallback to HTML
            html = self._to_html(title, include_graph, language)
            pdf_path = output_path.replace(".pdf", ".html")
            with open(pdf_path, "w", encoding="utf-8") as f:
                f.write(html)
            return pdf_path


# ============================================================================
# Convenience Functions
# ============================================================================

def summarize_documents(
    documents: List[Document],
    llm: Optional[BaseLLM] = None,
    language: Optional[str] = None,
) -> SummaryResult:
    """
    Convenience function to summarize documents.
    
    Args:
        documents: Documents to summarize.
        llm: Optional LLM instance.
        language: Output language.
        
    Returns:
        SummaryResult instance.
    """
    summarizer = Summarizer(llm=llm)
    return summarizer.summarize(documents, language=language)


def visualize_graph(
    graph: GraphStorage,
    format: str = "html",
    **kwargs,
) -> str:
    """
    Convenience function to visualize a knowledge graph.
    
    Args:
        graph: GraphStorage instance.
        format: Output format (html, dot, json, mermaid).
        **kwargs: Additional arguments for the visualizer.
        
    Returns:
        Formatted graph string.
    """
    visualizer = GraphVisualizer(graph)
    
    if format == "html":
        return visualizer.to_html(**kwargs)
    elif format == "dot":
        return visualizer.to_dot(**kwargs)
    elif format == "json":
        return visualizer.to_json(**kwargs)
    elif format == "mermaid":
        return visualizer.to_mermaid(**kwargs)
    else:
        raise ValueError(f"Unsupported format: {format}")


def export_data(
    output_path: str,
    format: str = "markdown",
    graph: Optional[GraphStorage] = None,
    **kwargs,
) -> str:
    """
    Convenience function to export data.
    
    Args:
        output_path: Output file path.
        format: Export format.
        graph: Optional GraphStorage.
        **kwargs: Additional arguments for the exporter.
        
    Returns:
        Output file path.
    """
    exporter = Exporter(graph=graph)
    return exporter.export(output_path, format=format, **kwargs)
